import json
import base64
import random
import os
import re
import time
import hmac

from  docx import  Document
from docx.shared import Pt, RGBColor
from  docx.oxml.ns import qn
from docx.shared import Inches

import smtplib
from email.mime.text import MIMEText
from email.header import Header


# 表单转字典
def formToDict(form):
    return json.loads(json.dumps(form))

# 生成token
def generate_token(key, expire = 3600):
    time_str = str(time.time() + expire)
    time_encode = time_str.encode("utf-8")

    sha1_hex = hmac.new(key.encode("utf-8"), time_encode, 'sha1').hexdigest()
    token = time_str + ":" + sha1_hex

    b64_token = base64.urlsafe_b64encode(token.encode("utf-8"))
    return b64_token.decode("utf-8")

def certify_token(key, token):
    token_str = base64.urlsafe_b64decode(token).decode("utf-8")
    token_list = token_str.split(":")

    if len(token_list)!=2:
        return False
    
    time_str = token_list[0]
    if float(time_str) < time.time():
        return False
    
    get_sha1 = token_list[1]
    cal_sha1 = hmac.new(key.encode("utf-8", time_str.encode("utf-8"), 'sha1')).hexdigest()

    if get_sha1!=cal_sha1:
        return False
    return True

# 生成验证码
def generateVerifyCode():
    return ''.join(random.sample(['1','2','3','4','5','6','7','8','9','0'], 6))

#发送验证码
def sendVerifyCode(email):
    sender = 'laizimengchris11@qq.com'
    pwd = 'xfuzywmxnbgpgaef'
    receiver = [email]

    code = generateVerifyCode()
    message = MIMEText(' SuperWrite 给您发送验证码: ' + code, 'plain', 'utf-8')
    message['From'] = Header('SuperWrite', 'utf-8')
    message['To'] = Header(email, 'utf-8')
    message['Subject'] = Header('SuperWrite 验证码', 'utf-8')

    try: 
        smtpObj = smtplib.SMTP_SSL("smtp.qq.com",465)
        smtpObj.login(sender, pwd)
        smtpObj.sendmail(sender, receiver, message.as_string())
        print('发送验证码到' + email + '成功')
        smtpObj.quit()
        return True, code
    
    except smtplib.SMTPException:
        print('发送邮件失败')
        smtpObj.quit()
        return False, _

#返回需要匹配的title
def pickTitles(data):
    titles = []
    for key in data:
        if re.search(r'title(\d+)$',key) != None:
            titles.append(data[key])
    return titles

# 处理指导书
def dealWithDocx(file, titles):
    result = []

    temp_title = ''
    real_title = ''
    temp_content = []

    #获取文档
    raw = Document(file)

    for para in raw.paragraphs:
        line = para.text
        print(line)
        is_match = False
        for title in titles:
            # 编码后匹配
            pattern_encode = title.encode('utf-8').strip()
            line_encode = line.encode('utf-8').strip()
            if re.search(pattern_encode, line_encode):
                print(title)
                if temp_title != '':
                    if temp_title == title:
                        break
                    result.append({
                        real_title : ''.join(temp_content)
                    })
                    temp_title = title
                    real_title = line
                    temp_content = []
                else:
                    temp_title = title
                    real_title = line
                    temp_content = []
                is_match = True
                break
        if not is_match:
            if temp_title != '':
                temp_content.append(line)
    result.append({
        real_title : ''.join(temp_content)
    })
    return result

def generateReport(textList, images):
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.size = Pt(10.5)
    document.styles['Normal'].font.color.rgb = RGBColor(0,0,0)
    for para in textList:
        print(para)
        if para['type'] == 'title':
            head = document.add_heading("", 2)
            run = head.add_run(para['value'])
            run.font.name=u'Cambria'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
            run.font.color.rgb = RGBColor(0,0,0)
        elif para['type'] == 'content':
            document.add_paragraph(para['value'])
    absPath = os.path.abspath('../send/test.docx')
    document.save('../send/test.docx')
    return absPath