import json
import base64
import random
import os
import re

from docx import Document
from docx.shared import Pt, RGBColor
from  docx.oxml.ns import qn
from docx.shared import Inches

import smtplib
from email.mime.text import MIMEText
from email.header import Header

#返回需要匹配的title
def pickTitles(data):
    titles = []
    for key in data:
        if re.search(r'title(\d+)$',key) != None:
            titles.append(data[key])
    return titles

# 处理指导书
def dealWithDocx(file, titles):
    result = []

    temp_title = ''
    real_title = ''
    temp_content = []

    #获取文档
    raw = Document(file)

    for para in raw.paragraphs:
        line = para.text
        print(line)
        is_match = False
        for title in titles:
            # 编码后匹配
            pattern_encode = title.encode('utf-8').strip()
            line_encode = line.encode('utf-8').strip()
            if re.search(pattern_encode, line_encode):
                print(title)
                if temp_title != '':
                    if temp_title == title:
                        break
                    result.append({
                        real_title : ''.join(temp_content)
                    })
                    temp_title = title
                    real_title = line
                    temp_content = []
                else:
                    temp_title = title
                    real_title = line
                    temp_content = []
                is_match = True
                break
        if not is_match:
            if temp_title != '':
                temp_content.append(line)
    result.append({
        real_title : ''.join(temp_content)
    })
    return result

def generateReport(textList, images):
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.size = Pt(10.5)
    document.styles['Normal'].font.color.rgb = RGBColor(0,0,0)
    for para in textList:
        print(para)
        if para['type'] == 'title':
            head = document.add_heading("", 2)
            run = head.add_run(para['value'])
            run.font.name=u'Cambria'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
            run.font.color.rgb = RGBColor(0,0,0)
        elif para['type'] == 'content':
            document.add_paragraph(para['value'])
    absPath = os.path.abspath('../send/test.docx')
    document.save('../send/test.docx')
    return absPath